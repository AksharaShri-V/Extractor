import streamlit as st
import PyPDF2
import openai
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tiktoken
import os
import time

# Set up OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def num_tokens_from_string(string: str, encoding_name: str) -> int:
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens

def split_into_chunks(text, max_tokens=3000):
    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)
    chunks = []
    current_chunk = []
    current_chunk_tokens = 0

    for token in tokens:
        if current_chunk_tokens + 1 > max_tokens:
            chunks.append(encoding.decode(current_chunk))
            current_chunk = []
            current_chunk_tokens = 0
        current_chunk.append(token)
        current_chunk_tokens += 1

    if current_chunk:
        chunks.append(encoding.decode(current_chunk))

    return chunks

def process_chunk_with_openai(chunk, is_first_chunk=False):
    system_instruction = """
    You are an advanced AI assistant specialized in processing text from PDFs. Your tasks are:

   1. Identify main headings, subheadings, and side headings. Use the following format:
       MAIN HEADING: text
       SUBHEADING: text
       SIDE HEADING: text

    2. Extract and format normal text paragraphs completely. Ensure no sentences or words are left incomplete.

    3. Identify any listed data or enumerated information and preserve its format. Do not add additional bullet points or enumeration if they already exist.

    4. Detect any tabular data structures within the text. For each detected table:
       - If the data is already in a bullet point or list format, preserve that format exactly.
       - Convert each row of the table into a bullet point
       - Start each bullet point with the first column's value
       - Include all values from all columns in the bullet point
       - Ensure that the relationship between all columns is clearly expressed
       - Do not omit any data for brevity
       For example, if a table has columns "Name", "Age", "Occupation", and "Salary", a row might be converted to:
       • John Doe, aged 30, works as a software engineer and earns $75,000 annually.

    5. Maintain the original order and context of the document while processing.

    6. Do not use any special characters or symbols for formatting except for the bullet points (•) for table data.

    7. It is crucial that you process and include ALL content from the given chunk. Do not truncate or omit any information.

    Your goal is to extract and transform the document content completely, preserving all original information and structure, while ensuring that tabular data is presented as bullet points that clearly convey the relationships between all data points.

    If this is the first chunk of the document, start with 'DOCUMENT START:'. If it's the last chunk, end with 'DOCUMENT END:'.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": f"Process the following text chunk from a PDF, following the instructions given. {'This is the first chunk of the document.' if is_first_chunk else ''}\n\n{chunk}"}
        ]
    )
    return response.choices[0].message['content']

def generate_summary(processed_text, filename):
    system_instruction = f"""
    You are an AI assistant specialized in creating concise and easily understandable summaries. Your task is to create a brief summary of the given text from the file "{filename}". The summary should:

    1. Start with a generic introductory sentence about the document.
    2. Capture the main ideas and key points of the document in a concise manner.
    3. Highlight only the most significant findings or conclusions.
    4. Mention only the most important data or statistics, if present.
    5. Be written in simple, clear language that is easy for a general audience to understand.
    6. Be no longer than 300 words, including the introductory and concluding sentences.
    7. Use bullet points for clarity where appropriate.
    8. End with a generic concluding sentence about the document's overall significance or relevance.
    9. Identify 5-7 keywords from the document and mark them with asterisks (*keyword*).

    Your goal is to provide a summary that gives readers a quick overview of the document's core content, making it distinctly different from the full processed text.

    Format the summary exactly as follows:
    SUMMARY START
    [Introductory sentence]
    
    [Main summary content with keywords marked]
    
    [Concluding sentence]
    SUMMARY END
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": f"Create a concise and easy-to-understand summary of the following processed text:\n\n{processed_text}"}
        ]
    )
    
    summary = response.choices[0].message['content']
    
    # Ensure the summary is properly formatted
    if not summary.startswith("SUMMARY START"):
        summary = "SUMMARY START\n" + summary
    if not summary.endswith("SUMMARY END"):
        summary += "\nSUMMARY END"
    
    return summary

def create_word_document(content):
    doc = Document()
    
    # Define styles for different heading levels
    styles = doc.styles
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 2'].font.size = Pt(14)
    styles['Heading 3'].font.size = Pt(12)
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('MAIN HEADING:'):
            doc.add_paragraph(line.replace('MAIN HEADING:', '').strip(), style='Heading 1')
        elif line.startswith('SUBHEADING:'):
            doc.add_paragraph(line.replace('SUBHEADING:', '').strip(), style='Heading 2')
        elif line.startswith('SIDE HEADING:'):
            doc.add_paragraph(line.replace('SIDE HEADING:', '').strip(), style='Heading 3')
        elif line.startswith('•'):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_summary_document(summary, filename):
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Executive Summary of {filename}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Process summary content
    summary_parts = summary.split("SUMMARY START")
    if len(summary_parts) > 1:
        summary_content = summary_parts[1].split("SUMMARY END")[0].strip()
    else:
        summary_content = summary  # If SUMMARY START is not found, use the entire summary

    paragraphs = summary_content.split('\n\n')

    for para in paragraphs:
        p = doc.add_paragraph()
        parts = para.split('*')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:  # Every odd part (1, 3, 5, ...) should be bold
                run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("Smart Extract!")

    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

    if uploaded_file is not None:
        if st.button("Process PDF"):
            with st.spinner("Processing PDF and generating Word document... This may take a while."):
                try:
                    file_contents = uploaded_file.read()
                    text = extract_text_from_pdf(io.BytesIO(file_contents))
                    chunks = split_into_chunks(text)
                    processed_chunks = []

                    # Calculate delay based on file size
                    file_size = len(file_contents)
                    base_delay = max(0.5, min(3, file_size / 1000000))  # 0.5 to 3 seconds

                    loading_messages = [
                        "Reading the document... 📄",
                        "Recognizing patterns... 🧠",
                        "Organizing content... 🗒️",
                        "Finalizing... 🚀"
                    ]

                    progress_placeholder = st.empty()
                    message_placeholder = st.empty()

                    for i, chunk in enumerate(chunks):
                        progress = (i + 1) / (len(chunks) + 1)  # +1 for summary generation
                        progress_placeholder.progress(progress)
                        
                        message_index = min(int(progress * len(loading_messages)), len(loading_messages) - 1)
                        message_placeholder.info(loading_messages[message_index])
                        
                        processed_chunk = process_chunk_with_openai(chunk, is_first_chunk=(i==0))
                        processed_chunks.append(processed_chunk)
                        
                        time.sleep(base_delay * (1 - progress))  # Delay decreases as progress increases

                    processed_text = "\n".join(processed_chunks)
                    word_buffer = create_word_document(processed_text)

                    # Generate summary
                    message_placeholder.info("Generating summary... 📝")
                    original_filename = os.path.splitext(uploaded_file.name)[0]
                    summary = generate_summary(processed_text, original_filename)
                    summary_buffer = create_summary_document(summary, original_filename)

                    progress_placeholder.empty()
                    message_placeholder.success("Processing complete. You can now download the full document and the summary.")
                    
                    # Store the buffers in session state
                    st.session_state['word_buffer'] = word_buffer
                    st.session_state['summary_buffer'] = summary_buffer
                    st.session_state['original_filename'] = original_filename

                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

    # Always show download buttons if buffers are available
    if 'word_buffer' in st.session_state and 'summary_buffer' in st.session_state:
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download Full Document",
                data=st.session_state['word_buffer'],
                file_name=f"{st.session_state['original_filename']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                label="📥 Download Summary",
                data=st.session_state['summary_buffer'],
                file_name=f"{st.session_state['original_filename']}_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
